"""Generate the camper roster + daily-totals docx for Oregon Tour de Outback 2026."""
import csv
import os
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
CSV_IN = os.path.join(HERE, "..", "Campers _ TdO 2026 - RV Hookup.csv")
OUT = os.path.join(HERE, "camper-roster-2026.docx")
LOGO = os.path.join(HERE, "..", "images", "logo-brown-red.png")

BRAND_RED = RGBColor(0xCC, 0x00, 0x00)
DARK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT = RGBColor(0xAA, 0xAA, 0xAA)

DAYS = ["Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
DAY_SHORT = ["Wed", "Thu", "Fri", "Sat", "Sun"]
# Arrival views only cover Wed-Fri — Sat/Sun "arrivals" are either zero or
# split-registrations from a Fri camper, not new check-ins to plan around.
ARRIVAL_DAYS = DAYS[:3]
ARRIVAL_SHORT = DAY_SHORT[:3]

# ============================================================
# LOAD DATA
# ============================================================
campers = []
if not os.path.exists(CSV_IN):
    raise SystemExit(f"CSV not found at: {CSV_IN}")

with open(CSV_IN, newline='', encoding='utf-8') as f:
    reader = csv.DictReader(f)
    for row in reader:
        # Filter cancellations into separate list later
        notes = (row.get('Notes') or '').strip()
        c = {
            'last': row['Last Name'].strip(),
            'first': row['First Name'].strip(),
            'phone': row['Phone'].strip(),
            'email': row['Email'].strip(),
            'type': row['H/U or Dry?'].strip(),
            'nights': row['# Nights'].strip(),
            'days': {d: (row.get(d, '').strip() == '1') for d in DAYS},
            'notes': notes,
            'cancelled': 'cancel' in notes.lower(),
        }
        campers.append(c)

active = [c for c in campers if not c['cancelled']]
cancelled = [c for c in campers if c['cancelled']]
active.sort(key=lambda c: (c['last'].lower(), c['first'].lower()))

# Type display
def type_label(t):
    if 'Dry' in t and 'H/U' in t:
        return 'Dry+HU'
    if 'H/U' in t:
        return 'Hookup'
    return 'Dry'

# Compute arrivals per day (the first night a camper is on-site)
arrivals_by_day = defaultdict(list)
arrival_totals = {d: {'Dry': 0, 'Hookup': 0, 'Both': 0} for d in DAYS}
for c in active:
    first_day = next((d for d in DAYS if c['days'][d]), None)
    if first_day:
        arrivals_by_day[first_day].append(c)
        t = type_label(c['type'])
        bucket = 'Both' if t == 'Dry+HU' else ('Hookup' if t == 'Hookup' else 'Dry')
        arrival_totals[first_day][bucket] += 1


# ============================================================
# DOCX BUILDING HELPERS
# ============================================================
def set_cell_borders(cell, color='CCCCCC'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        b.set(qn('w:space'), '0')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def table_borders(table, color='BBBBBB'):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        b.set(qn('w:space'), '0')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def write(cell, text, *, bold=False, italic=False, size=9, color=DARK, align=None):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r


def section_header(doc, text, *, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = BRAND_RED
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:color'), 'CC0000')
    bot.set(qn('w:space'), '1')
    pBdr.append(bot)
    pPr.append(pBdr)


# ============================================================
# DOCUMENT SETUP — LANDSCAPE
# ============================================================
doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = Inches(0.45)
section.bottom_margin = Inches(0.45)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# Header: logo right, title left
header = section.header
htbl = header.add_table(rows=1, cols=2, width=Inches(10))
htbl.autofit = False
htbl.columns[0].width = Inches(8.5)
htbl.columns[1].width = Inches(1.5)
left, right = htbl.cell(0, 0), htbl.cell(0, 1)
write(left, "OREGON TOUR DE OUTBACK 2026", bold=True, size=11, color=DARK)
p2 = left.add_paragraph()
p2.paragraph_format.space_after = Pt(0)
r2 = p2.add_run("Camper Roster — Lake County Fairgrounds")
r2.italic = True
r2.font.size = Pt(9)
r2.font.color.rgb = GRAY
right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
right.paragraphs[0].add_run().add_picture(LOGO, width=Inches(0.85))

# ============================================================
# TITLE
# ============================================================
t = doc.add_paragraph()
tr = t.add_run("Camper Roster")
tr.bold = True
tr.font.size = Pt(18)
tr.font.color.rgb = DARK
t.paragraph_format.space_after = Pt(0)

s = doc.add_paragraph()
sr = s.add_run("Wednesday June 24 – Sunday June 28, 2026  ·  1900 N 4th St, Lakeview, OR")
sr.italic = True
sr.font.size = Pt(10)
sr.font.color.rgb = GRAY
s.paragraph_format.space_after = Pt(2)

# ============================================================
# AT-A-GLANCE ARRIVALS
# ============================================================
section_header(doc, "Arrivals at a Glance")
n_arrival_cols = 1 + len(ARRIVAL_DAYS)
tot_table = doc.add_table(rows=4, cols=n_arrival_cols)
tot_table.alignment = WD_TABLE_ALIGNMENT.LEFT
tot_table.autofit = False
widths = [Inches(1.4)] + [Inches(1.4)] * len(ARRIVAL_DAYS)
for i, w in enumerate(widths):
    tot_table.columns[i].width = w

# header row
write(tot_table.cell(0, 0), "", bold=True, size=9.5)
shade_cell(tot_table.cell(0, 0), 'F2F2F2')
for i, d in enumerate(ARRIVAL_SHORT):
    c = tot_table.cell(0, i + 1)
    write(c, d, bold=True, size=9.5, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(c, 'F2F2F2')

# Dry row
write(tot_table.cell(1, 0), "Dry Camping", bold=True, size=9.5)
for i, d in enumerate(ARRIVAL_DAYS):
    v = arrival_totals[d]['Dry'] + arrival_totals[d]['Both']
    write(tot_table.cell(1, i + 1), str(v), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

# Hookup row
write(tot_table.cell(2, 0), "RV Hookup", bold=True, size=9.5)
for i, d in enumerate(ARRIVAL_DAYS):
    v = arrival_totals[d]['Hookup'] + arrival_totals[d]['Both']
    write(tot_table.cell(2, i + 1), str(v), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

# Total row
write(tot_table.cell(3, 0), "Total Arriving", bold=True, size=10, color=BRAND_RED)
shade_cell(tot_table.cell(3, 0), 'FAF2F2')
for i, d in enumerate(ARRIVAL_DAYS):
    v = arrival_totals[d]['Dry'] + arrival_totals[d]['Hookup'] + arrival_totals[d]['Both']
    c = tot_table.cell(3, i + 1)
    write(c, str(v), bold=True, size=11, color=BRAND_RED, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(c, 'FAF2F2')

table_borders(tot_table)

note = doc.add_paragraph()
nr = note.add_run("Each camper is counted on the day they first arrive. Saturday and Sunday arrivals are not shown — there are no new check-ins after Friday for the 2026 event.")
nr.italic = True
nr.font.size = Pt(8.5)
nr.font.color.rgb = GRAY
note.paragraph_format.space_after = Pt(2)

# ============================================================
# ARRIVALS BY DAY
# ============================================================
section_header(doc, "Arrivals by Day")
intro = doc.add_paragraph()
ir = intro.add_run("Campers grouped by the day they first arrive. Use this to know who to expect at check-in.")
ir.italic = True
ir.font.size = Pt(8.5)
ir.font.color.rgb = GRAY
intro.paragraph_format.space_after = Pt(3)

n_arr_cols = len(ARRIVAL_DAYS)
arr_table = doc.add_table(rows=2, cols=n_arr_cols)
arr_table.alignment = WD_TABLE_ALIGNMENT.LEFT
arr_table.autofit = False
for i in range(n_arr_cols):
    arr_table.columns[i].width = Inches(3.2)

for i, d in enumerate(ARRIVAL_SHORT):
    full_day = ARRIVAL_DAYS[i]
    arr = arrivals_by_day.get(full_day, [])
    header_cell = arr_table.cell(0, i)
    body_cell = arr_table.cell(1, i)

    # header with count
    header_cell.paragraphs[0].clear()
    p = header_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{d}  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = DARK
    r2 = p.add_run(f"({len(arr)} arriving)")
    r2.italic = True
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = BRAND_RED
    shade_cell(header_cell, 'F2F2F2')

    body_cell.paragraphs[0].clear()
    if not arr:
        write(body_cell, "—", size=9, color=LIGHT)
    else:
        for j, c in enumerate(arr):
            p = body_cell.paragraphs[0] if j == 0 else body_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"{c['last']}, {c['first']}")
            r.font.size = Pt(8.5)
            tr = p.add_run(f"  ({type_label(c['type'])})")
            tr.italic = True
            tr.font.size = Pt(8)
            tr.font.color.rgb = GRAY

table_borders(arr_table)

# Page break before master roster
doc.add_page_break()

# ============================================================
# MASTER ROSTER
# ============================================================
section_header(doc, f"Master Roster  ·  {len(active)} confirmed campers")

# Columns: # | Last, First | Type | Phone | Email | W | T | F | S | Su | # Nights | Notes
col_widths = [
    Inches(0.3),   # #
    Inches(1.6),   # name
    Inches(0.55),  # type
    Inches(0.95),  # phone
    Inches(1.8),   # email
    Inches(0.35), Inches(0.35), Inches(0.35), Inches(0.35), Inches(0.35),  # days
    Inches(0.5),   # nights
    Inches(1.8),   # notes
]

n_cols = len(col_widths)
roster_rows = len(active) + 12  # + write-in rows
tbl = doc.add_table(rows=1 + roster_rows, cols=n_cols)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.autofit = False
for i, w in enumerate(col_widths):
    tbl.columns[i].width = w

# Header row
headers = ["#", "Camper (Last, First)", "Type", "Phone", "Email",
           "W", "T", "F", "S", "Su", "Nts", "Notes"]
for i, h in enumerate(headers):
    c = tbl.cell(0, i)
    c.width = col_widths[i]
    align = WD_ALIGN_PARAGRAPH.CENTER if i not in (1, 3, 4, 11) else WD_ALIGN_PARAGRAPH.LEFT
    write(c, h, bold=True, size=9, color=DARK, align=align)
    shade_cell(c, 'F2F2F2')

# Active campers
for idx, camper in enumerate(active, start=1):
    row = tbl.rows[idx]
    for i, w in enumerate(col_widths):
        row.cells[i].width = w
    # #
    write(row.cells[0], str(idx), size=8.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    # name
    write(row.cells[1], f"{camper['last']}, {camper['first']}", size=9)
    # type
    t_lbl = type_label(camper['type'])
    color = BRAND_RED if 'HU' in t_lbl or 'Hook' in t_lbl else DARK
    write(row.cells[2], t_lbl, size=8.5, color=color, align=WD_ALIGN_PARAGRAPH.CENTER)
    # phone
    write(row.cells[3], camper['phone'], size=8.5)
    # email
    write(row.cells[4], camper['email'], size=8)
    # day checkmarks
    for di, d in enumerate(DAYS):
        cell = row.cells[5 + di]
        mark = '✔' if camper['days'][d] else ''
        write(cell, mark, size=11, color=BRAND_RED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    # nights
    write(row.cells[10], camper['nights'], size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    # notes
    write(row.cells[11], camper['notes'], size=8, italic=True, color=GRAY)

# Write-in rows
writein_start_idx = len(active) + 1
for i in range(12):
    row = tbl.rows[writein_start_idx + i]
    for ci, w in enumerate(col_widths):
        row.cells[ci].width = w
    # leave cells empty — just blank rows
    # Shade the row number column lightly so it reads "fill me in"
    write(row.cells[0], str(len(active) + 1 + i), size=8.5, color=LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER)

table_borders(tbl)

# Legend
legend = doc.add_paragraph()
legend.paragraph_format.space_before = Pt(6)
legend.paragraph_format.space_after = Pt(0)
lr = legend.add_run("Legend:  ")
lr.italic = True
lr.font.size = Pt(8.5)
lr.font.color.rgb = GRAY
parts = [
    ("Type: ", DARK, False),
    ("Dry", DARK, True),
    (" = dry camping  ·  ", GRAY, False),
    ("Hookup", BRAND_RED, True),
    (" = full RV hookup  ·  ", GRAY, False),
    ("Dry+HU", DARK, True),
    (" = mixed nights.  Days: ", GRAY, False),
    ("W = Wed", DARK, True),
    (", T = Thu, F = Fri, S = Sat, Su = Sun.  ", GRAY, False),
    ("✔", BRAND_RED, True),
    (" = staying that night.", GRAY, False),
]
for text, color, bold in parts:
    r = legend.add_run(text)
    r.font.size = Pt(8.5)
    r.font.color.rgb = color
    r.bold = bold

# ============================================================
# CANCELLATIONS
# ============================================================
if cancelled:
    section_header(doc, f"Cancellations  ·  {len(cancelled)}", size=11)
    for c in cancelled:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{c['last']}, {c['first']}  —  {c['phone']}  —  {c['email']}")
        r.font.size = Pt(9)
        r.font.color.rgb = LIGHT

# ============================================================
# DATA-QUALITY FLAGS (only if duplicates exist)
# ============================================================
seen_phone = defaultdict(list)
for c in active:
    if c['phone']:
        seen_phone[c['phone']].append(f"{c['last']}, {c['first']}")
duplicates = {k: v for k, v in seen_phone.items() if len(v) > 1}
if duplicates:
    section_header(doc, "Possible Duplicates — Please Review", size=11)
    for phone, names in duplicates.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"Same phone ({phone}): {' / '.join(names)}")
        r.font.size = Pt(9)
        r.italic = True
        r.font.color.rgb = GRAY

# ============================================================
# FOOTER NOTE
# ============================================================
foot = doc.add_paragraph()
foot.paragraph_format.space_before = Pt(10)
fr = foot.add_run(
    "Generated from BikeReg camping registrations. Roster is a working draft — "
    "verify with day-of leads at check-in. Update this document as new registrations come in."
)
fr.italic = True
fr.font.size = Pt(8)
fr.font.color.rgb = GRAY

doc.save(OUT)
print(f"Wrote: {OUT}")
print(f"Active campers: {len(active)}  ·  Cancellations: {len(cancelled)}")
print("Arrivals by day (Wed-Fri only shown in doc):")
for d in DAYS:
    t = arrival_totals[d]
    total = t['Dry'] + t['Hookup'] + t['Both']
    marker = "  " if d in ARRIVAL_DAYS else " *"
    print(f"  {d}:{marker}{total} arriving (Dry {t['Dry'] + t['Both']}, Hookup {t['Hookup'] + t['Both']})")
