"""Generate volunteer schedule + directory docx for Oregon Tour de Outback 2026."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "volunteer-schedule-2026.docx")
LOGO = os.path.join(os.path.dirname(__file__), "..", "images", "logo-brown-red.png")

BRAND_RED = RGBColor(0xCC, 0x00, 0x00)
DARK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# Narrow margins for printability
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# Logo in the header (right-aligned, small, won't disturb body flow)
header = doc.sections[0].header
htbl = header.add_table(rows=1, cols=2, width=Inches(7.3))
htbl.autofit = False
htbl.columns[0].width = Inches(5.5)
htbl.columns[1].width = Inches(1.8)
left_cell = htbl.cell(0, 0)
right_cell = htbl.cell(0, 1)
lp = left_cell.paragraphs[0]
lr = lp.add_run("OREGON TOUR DE OUTBACK 2026")
lr.bold = True
lr.font.size = Pt(11)
lr.font.color.rgb = DARK
lp2 = left_cell.add_paragraph()
lr2 = lp2.add_run("Volunteer Schedule & Directory")
lr2.italic = True
lr2.font.size = Pt(9)
lr2.font.color.rgb = GRAY
rp = right_cell.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rp.add_run().add_picture(LOGO, width=Inches(0.85))


def add_day_header(doc, title, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BRAND_RED
    # thin bottom border on the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:color'), 'CC0000')
    bot.set(qn('w:space'), '1')
    pBdr.append(bot)
    pPr.append(pBdr)
    if note:
        np = doc.add_paragraph()
        np.paragraph_format.space_after = Pt(4)
        nr = np.add_run(note)
        nr.italic = True
        nr.font.size = Pt(9)
        nr.font.color.rgb = GRAY


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def add_schedule_table(doc, rows):
    """rows: list of (time, event_lines, volunteers).
    event_lines may be a string or list of strings (sub-items)."""
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Inches(1.1), Inches(4.6), Inches(1.6)]
    for col_idx, w in enumerate(widths):
        table.columns[col_idx].width = w

    # header row
    hdr = table.rows[0]
    for i, label in enumerate(("Time", "Event / Need", "Volunteer(s)")):
        c = hdr.cells[i]
        c.width = widths[i]
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK
        shade_cell(c, 'F2F2F2')

    # body rows
    for r_idx, (time, event, vol) in enumerate(rows, start=1):
        row = table.rows[r_idx]
        for i, w in enumerate(widths):
            row.cells[i].width = w

        # time
        tcell = row.cells[0]
        tcell.paragraphs[0].clear()
        tr = tcell.paragraphs[0].add_run(time)
        tr.bold = True
        tr.font.size = Pt(9.5)
        tr.font.color.rgb = BRAND_RED

        # event (may be multi-line)
        ecell = row.cells[1]
        ecell.paragraphs[0].clear()
        lines = event if isinstance(event, list) else [event]
        for li, line in enumerate(lines):
            p = ecell.paragraphs[0] if li == 0 else ecell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            er = p.add_run(line)
            er.font.size = Pt(9.5)

        # volunteers
        vcell = row.cells[2]
        vcell.paragraphs[0].clear()
        vlines = vol if isinstance(vol, list) else [vol]
        for li, line in enumerate(vlines):
            p = vcell.paragraphs[0] if li == 0 else vcell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            vr = p.add_run(line)
            vr.font.size = Pt(9.5)
            vr.italic = True

    # apply thin borders to the table
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'CCCCCC')
        b.set(qn('w:space'), '0')
        tblBorders.append(b)
    tblPr.append(tblBorders)


# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("Volunteer Schedule")
tr.bold = True
tr.font.size = Pt(18)
tr.font.color.rgb = DARK
title.paragraph_format.space_after = Pt(0)

sub = doc.add_paragraph()
sr = sub.add_run("Prep Week through Ride Day  ·  June 23–27, 2026  ·  Lake County Fairgrounds, Lakeview, OR")
sr.italic = True
sr.font.size = Pt(10)
sr.font.color.rgb = GRAY
sub.paragraph_format.space_after = Pt(6)

# ============================================================
# TUESDAY
# ============================================================
add_day_header(doc, "TUESDAY  ·  June 23")
add_schedule_table(doc, [
    ("9:00 AM",
     ["Signage: pick up signage at Shed Quarters (wooden, plastic ones)",
      "Attach to boards, make paper signs, directional signs",
      "Placement of Gravel Route signage"],
     "Donna, Lesa, Marie"),
    ("10:00 AM",
     "Confirm Safeway order pick-up for Thursday (AM or PM??)",
     "Marie"),
])

# ============================================================
# WEDNESDAY
# ============================================================
add_day_header(doc, "WEDNESDAY  ·  June 24", note="Waivers needed for Wednesday volunteers.")
add_schedule_table(doc, [
    ("12:00 PM",
     ["Tent set-up — check to Dayton Alves $500",
      "Stage placement, arrange electrical for band(s)"],
     ["Dayton Alves, Fairgrounds helper",
      "Marty C., Donna, Marie, David"]),
])

# ============================================================
# THURSDAY
# ============================================================
add_day_header(doc, "THURSDAY  ·  June 25", note="Waivers needed for new volunteers.")
add_schedule_table(doc, [
    ("8:00 – 10:00 AM", "Spray paint on gravel route", "Donna"),
    ("12:00 PM", "Mt Mazama water delivered", "Mt Mazama, Marie"),
    ("12:00 PM", "Porta-potties delivered and confirm sites", "Lesa"),
    ("12:00 PM",
     ["Tables and chairs set up; place tablecloths on each table, tape down each side",
      "Tables (2–3) for Pizza — label",
      "Tables (2) for Registration — signage labeled"],
     ["Saul, David, Donna, Hugh, Marie (fairgrounds help), Marie",
      "",
      "David"]),
    ("2:00 PM", "Clean bathrooms and showers", "Sara Bowersox"),
    ("2:00 PM?", "Band set up", "Marty C."),
    ("3:00 PM", "RV trailers to RV park (Tucker's and Griffith's)", "Tim, Anneke"),
    ("4:00 PM",
     ["Shed Quarters (alleyway of 839 South G St)",
      "Load trailer & unload at fairgrounds"],
     "Hugh, Lesa, SAR helper, David, Saul, Tim, Marie, Donna"),
    ("5:00 PM",
     ["Organize headquarters items: registration area, signage",
      "Kitchen — get organized for Aide Station help"],
     ["David, Donna, Lesa",
      "Marie"]),
])

# ============================================================
# FRIDAY
# ============================================================
add_day_header(doc, "FRIDAY  ·  June 26  —  Prep & Rider Check-In",
               note="Waivers for any new volunteers.")
add_schedule_table(doc, [
    ("7:00 AM", "Contact police *677 — advise of cyclists on road", "Marie"),
    ("7:00 AM", "Pick up eggs, potatoes", "Marie"),
    ("8:00 AM",
     ["Aide Station prep — wash and dry produce",
      "Sort fruit in portions and assemble cooler (as much as possible)",
      "Assemble utensils and place in collapsible boxes",
      "Label with inventory of cooler; label road or gravel ride",
      "Assemble SAG coolers — Gatorade, water",
      "Pop-up tents labeled if need be",
      "Signage up on reader board",
      "Signage on Headquarters",
      "Feather flags placed",
      "Start Line placed, generator",
      "Arrange Packet Pick-Up",
      "Event board",
      "Map board",
      "Cash box & quarters"],
     ["Danné, Shar, Kit",
      "",
      "",
      "Marie",
      "",
      "Lesa",
      "Brianna",
      "David",
      "David",
      "David, Tim",
      "David, Callie?",
      "Lesa, Donna",
      "Lesa",
      "Marie"]),
    ("1:00 PM",
     ["Cyclists check in, packet pick-up; possible late registrations",
      "Routes confirmed, wristbands secured, raffle tickets, drink tickets"],
     ["Jeanne, Alan (Donna, Lesa, Marie rotating)",
      "David"]),
    ("3:00 PM", "Shake-Out Ride", "Donna, Chris"),
    ("4:00 PM", "Mobile Bar (4–9 PM)", "Jennifer Simpson"),
    ("5:00 – 5:15 PM", "Pizza volunteers — instructions on serving, gloves", "Claire, Ann, Susan"),
    ("5:30 PM", "Meet & Greet Pizza Party (5:30 – 7:00 PM)", "Downtown Bakery, Committee"),
    ("7:00 PM",
     ["Announcements — welcome; check back in (we are in the \"outback\")",
      "Raffle",
      "Announce next year date — June 26th"],
     ["Donna, Marie, David",
      "Cyclist — Wanda",
      "David"]),
])

# Force page break before Ride Day for clean print
doc.add_page_break()

# ============================================================
# SATURDAY — RIDE DAY
# ============================================================
add_day_header(doc, "SATURDAY  ·  June 27  —  RIDE DAY")
add_schedule_table(doc, [
    ("5:00 AM", "Make all coffee, locate cups, cut bananas, eggs, fill urns", "Lesa, Marie"),
    ("5:30 AM", "Complete Aide Station coolers — food from fridge distributed to coolers", "Marie, Shar"),
    ("6:00 AM", "SAR — check out radios, coolers, ice, tents, water, beverage containers", "Lesa"),
    ("6:15 AM",
     ["SAR refresher — go over food display on tables at Aide Station",
      "Wristbands for volunteer meal; ask volunteers to take photos of riders at their Aide Station",
      "Road Route signage check"],
     "Lesa, SAR"),
    ("6:15 AM", "Registration (LATE) set-up", "Jeanne, Alan, Callie"),
    ("6:15 AM", "Pastries & Photos — display food, eggs, coffee cups, coffee", "Rotary, Callie"),
    ("6:30 AM", "Route ??? — rotating helpers", "Donna, Lesa, Marie"),
    ("6:30 AM",
     ["SAG Meeting — signage on sweep",
      "Don stays at Willow Bridge until all cyclists are over Hwy 395"],
     "Lesa, Tim, Don"),
    ("6:30 AM",
     ["\"Road Helpers\" — guide / stop traffic for cyclists:",
      "  Center St. & Hwy 395",
      "  Hollow Bridge Rd. & Hwy 395"],
     ["",
      "Mark Smith",
      "Don Brown*"]),
    ("6:45 AM", "Announcements — welcome; check-in due to \"outback\"; course closes at 3 PM", "Donna, Marie"),
    ("6:50 AM", "Photographers leave", "Donna, David, Chris"),
    ("7:00 AM", "Ride Start — Craig Price (police vehicle) leads cyclists across the road; cowbell ringers", "Committee"),
    ("7:30 AM", "SAG drivers leave", "Tim, Don"),
    ("8:00 AM", "Marie leaves Headquarters and drives to 1st road Aide Station", "Marie"),
    ("8:00 AM", "Band set up — Dan Decker and Missing Identity", "Callie"),
    ("9:00 AM", "Check trash", "Callie"),
    ("10:00 AM", "Road Route TOWN signage down", "Marie"),
    ("11:00 AM", "Showers — $35 in quarters to each shower in a dish with signage", "Callie"),
    ("11:00 AM", "Lunch begins (11 – 3 PM from food truck; after that boxed in cooler)", "Lee Salvidor"),
    ("11:00 AM", "Clean kitchen, utensils, sinks, fridge, etc.", "Marie, Callie"),
    ("12:00 PM", "Live Music — Missing Identity until 4 PM", "Dan Decker"),
    ("12:00 PM", "Trash check, lunch help, greeting cyclists, check showers for change", "Callie"),
    ("12:00 PM", "School Showers open", "Lesa"),
    ("12:00 PM", "Massage Therapist starts", ""),
    ("1:00 – 2:00 PM", "Committee back to Headquarters, meet cyclists", "Committee"),
    ("3:00 PM", "Route closes — signs packed up, Aide Stations close, final sweep", "SAR, SAG"),
    ("3:00 PM", "Lunch boxed up in coolers", "Lee Salvidor"),
    ("4:00 – 5:00 PM", "Load trailer with as much as possible without disturbing party", "Committee"),
])

# ============================================================
# DIRECTORY
# ============================================================
add_day_header(doc, "VOLUNTEER DIRECTORY")
np = doc.add_paragraph()
np.paragraph_format.space_after = Pt(4)
nr = np.add_run("Contact info as of 2026-06-14. Update this section as needed.")
nr.italic = True
nr.font.size = Pt(9)
nr.font.color.rgb = GRAY

directory = [
    ("First Name", "Last Name", "Role", "Phone", "Email"),
    ("Seth", "Ballaine", "", "971-400-0330", "seth.ballaine@gmail.com"),
    ("Don", "Brown", "SAG Driver", "541-840-3983", "donnwb@gmail.com"),
    ("Sharylinn", "McClain", "AID Stations Food", "541-219-0892", ""),
    ("Kit", "Thornton", "AID Stations Food", "541-417-0859", ""),
    ("Callie", "Elliot", "", "541-521-2696", ""),
    ("Mark", "Smith", "", "505-710-4380", ""),
    ("Hugh", "Cahill", "", "", ""),
    ("Danné", "Barry", "", "541-219-2219", ""),
    ("Betsy", "Miller", "", "", ""),
    ("Tim", "Tucker", "", "541-417-0871", ""),
    ("Marty", "Chaloupka", "", "971-506-6072", ""),
    ("Chris", "Stanton", "", "", ""),
    ("Susan", "Albertson", "", "541-417-0978", ""),
    ("Claire", "Reaume", "", "541-219-2334", ""),
    ("Ann", "Sabin", "", "541-219-0135", ""),
]

dtable = doc.add_table(rows=len(directory), cols=5)
dtable.alignment = WD_TABLE_ALIGNMENT.LEFT
dtable.autofit = False
dwidths = [Inches(0.9), Inches(1.1), Inches(1.7), Inches(1.3), Inches(2.3)]
for ci, w in enumerate(dwidths):
    dtable.columns[ci].width = w

for ri, row_data in enumerate(directory):
    row = dtable.rows[ri]
    for ci, value in enumerate(row_data):
        cell = row.cells[ci]
        cell.width = dwidths[ci]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(value)
        run.font.size = Pt(9.5)
        if ri == 0:
            run.bold = True
            run.font.color.rgb = DARK
            shade_cell(cell, 'F2F2F2')

# borders on directory
tbl = dtable._tbl
tblPr = tbl.tblPr
tblBorders = OxmlElement('w:tblBorders')
for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
    b = OxmlElement(f'w:{edge}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '4')
    b.set(qn('w:color'), 'CCCCCC')
    b.set(qn('w:space'), '0')
    tblBorders.append(b)
tblPr.append(tblBorders)

# Footer note
foot = doc.add_paragraph()
foot.paragraph_format.space_before = Pt(10)
fr = foot.add_run("Schedule and assignments are working drafts — confirm with day-of leads before relying on a specific time or person.")
fr.italic = True
fr.font.size = Pt(8.5)
fr.font.color.rgb = GRAY

doc.save(OUT)
print(f"Wrote: {OUT}")
